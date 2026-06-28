"""Tests for the brief-upload feature: deterministic text extraction, the
LLM adapter (mocked — no real API calls), and the two endpoints.

Run from the project root::

    python -m pytest tests/test_brief.py
    # or directly:
    python tests/test_brief.py
"""

import io
import os
import sys
import types

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, ROOT)
sys.path.insert(0, os.path.join(ROOT, "src"))

from fastapi.testclient import TestClient  # noqa: E402

from src.api import brief_extract  # noqa: E402
from src.api import brief_parser  # noqa: E402
from src.api.brief_parser import (  # noqa: E402
    BriefParserError,
    BriefParserUnavailable,
    DraftTask,
    _ModelOutput,
    _reconcile_skills,
    parse_brief,
)
from src.api.app import app  # noqa: E402

client = TestClient(app)


# ---------------------------------------------------------------------------
# Fixtures: build real .docx / .pdf bytes in memory (nothing is persisted).
# ---------------------------------------------------------------------------

def _make_docx(paragraphs, table_rows=None) -> bytes:
    import docx

    document = docx.Document()
    for p in paragraphs:
        document.add_paragraph(p)
    if table_rows:
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for r, row in enumerate(table_rows):
            for c, val in enumerate(row):
                table.cell(r, c).text = val
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def _make_text_pdf(text: str) -> bytes:
    """Build a minimal, valid, text-based PDF with a correct xref table."""
    safe = (
        text.replace("\\", r"\\").replace("(", r"\(").replace(")", r"\)")
    ).encode("latin-1")
    stream = b"BT /F1 24 Tf 72 720 Td (" + safe + b") Tj ET"
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>",
        b"<< /Length " + str(len(stream)).encode() + b" >>\nstream\n" + stream
        + b"\nendstream",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    out = bytearray(b"%PDF-1.4\n")
    offsets = []
    for i, obj in enumerate(objects, start=1):
        offsets.append(len(out))
        out += str(i).encode() + b" 0 obj\n" + obj + b"\nendobj\n"
    xref_pos = len(out)
    n = len(objects) + 1
    out += b"xref\n0 " + str(n).encode() + b"\n"
    out += b"0000000000 65535 f \n"
    for off in offsets:
        out += ("%010d 00000 n \n" % off).encode()
    out += b"trailer\n<< /Size " + str(n).encode() + b" /Root 1 0 R >>\n"
    out += b"startxref\n" + str(xref_pos).encode() + b"\n%%EOF"
    return bytes(out)


# ---------------------------------------------------------------------------
# Extraction (deterministic, no LLM)
# ---------------------------------------------------------------------------

def test_extract_docx_returns_text():
    content = _make_docx(["Build a mobile onboarding flow.", "Add analytics."])
    result = brief_extract.extract_brief_text(content, "brief.docx")
    assert result.file_type == "docx"
    assert "onboarding" in result.text
    assert "analytics" in result.text.lower()
    assert result.char_count > 0
    assert result.truncated is False


def test_extract_docx_includes_table_cells():
    content = _make_docx(
        ["Project brief"],
        table_rows=[["Task", "Owner"], ["Design login screen", "TBD"]],
    )
    result = brief_extract.extract_brief_text(content, "brief.docx")
    assert "Design login screen" in result.text


def test_extract_pdf_returns_text():
    content = _make_text_pdf("Design the onboarding experience for new users")
    result = brief_extract.extract_brief_text(content, "brief.pdf")
    assert result.file_type == "pdf"
    # Tolerate pypdf spacing differences by comparing space-stripped text.
    assert "onboarding" in result.text.replace(" ", "").lower()


def test_extract_scanned_pdf_is_rejected():
    # A valid PDF with no extractable text simulates a scanned/image-only file.
    content = _make_text_pdf("")
    try:
        brief_extract.extract_brief_text(content, "scan.pdf")
        assert False, "expected BriefExtractionError"
    except brief_extract.BriefExtractionError as exc:
        assert exc.status == 422
        assert "scanned" in exc.message.lower()


def test_extract_rejects_unsupported_extension():
    try:
        brief_extract.extract_brief_text(b"hello world", "notes.txt")
        assert False
    except brief_extract.BriefExtractionError as exc:
        assert exc.status == 415


def test_extract_rejects_old_doc():
    try:
        brief_extract.extract_brief_text(b"\xd0\xcf\x11\xe0data", "old.doc")
        assert False
    except brief_extract.BriefExtractionError as exc:
        assert exc.status == 415


def test_extract_rejects_empty():
    try:
        brief_extract.extract_brief_text(b"", "brief.docx")
        assert False
    except brief_extract.BriefExtractionError as exc:
        assert exc.status == 422


def test_extract_rejects_oversized():
    big = b"x" * (brief_extract.MAX_BYTES + 1)
    try:
        brief_extract.extract_brief_text(big, "brief.docx")
        assert False
    except brief_extract.BriefExtractionError as exc:
        assert exc.status == 413


def test_extract_rejects_corrupt_docx():
    try:
        brief_extract.extract_brief_text(b"not a real docx file", "brief.docx")
        assert False
    except brief_extract.BriefExtractionError as exc:
        assert exc.status == 422


# ---------------------------------------------------------------------------
# Skill reconciliation (deterministic — never trust the model blindly)
# ---------------------------------------------------------------------------

def test_reconcile_flags_unmatched_skill():
    tasks = [
        DraftTask(task="A", required_skill="UX", effort_hours=10),
        DraftTask(task="B", required_skill="Blockchain", effort_hours=5),
    ]
    unmatched = _reconcile_skills(tasks, ["UX", "Research", "Python"])
    assert unmatched == ["Blockchain"]
    assert tasks[0].needs_user_review is False
    assert tasks[1].needs_user_review is True
    assert "Blockchain" in (tasks[1].review_reason or "")


def test_reconcile_normalises_casing():
    tasks = [DraftTask(task="A", required_skill="ux", effort_hours=10)]
    _reconcile_skills(tasks, ["UX", "Research"])
    assert tasks[0].required_skill == "UX"
    assert tasks[0].needs_user_review is False


def test_reconcile_repairs_nonpositive_effort():
    tasks = [DraftTask(task="A", required_skill="UX", effort_hours=0)]
    _reconcile_skills(tasks, ["UX"])
    assert tasks[0].effort_hours > 0
    assert tasks[0].effort_is_estimated is True


# ---------------------------------------------------------------------------
# Parser adapter (Anthropic SDK mocked — no network)
# ---------------------------------------------------------------------------

class _FakeMessages:
    def __init__(self, response, capture=None):
        self._response = response
        self._capture = capture

    def parse(self, **kwargs):
        # Record the exact kwargs so tests can assert the real call shape
        # (model name, output_format present, no sampling params).
        if self._capture is not None:
            self._capture.clear()
            self._capture.update(kwargs)
        if isinstance(self._response, Exception):
            raise self._response
        return self._response


class _FakeResponse:
    def __init__(self, parsed, stop_reason="end_turn"):
        self.parsed_output = parsed
        self.stop_reason = stop_reason


def _install_fake_anthropic(response, capture=None):
    """Inject a fake ``anthropic`` module; return it so tests can build errors.

    The fake mirrors the real call shape: ``Anthropic().messages.parse(...)``
    returning an object with ``parsed_output`` + ``stop_reason`` (the real SDK
    returns ``ParsedMessage``, a ``Message`` subclass with both), and raising
    ``pydantic.ValidationError`` / typed errors just as the SDK does.
    """
    mod = types.ModuleType("anthropic")

    class RateLimitError(Exception):
        pass

    class AuthenticationError(Exception):
        pass

    class BadRequestError(Exception):
        pass

    class APIConnectionError(Exception):
        pass

    mod.RateLimitError = RateLimitError
    mod.AuthenticationError = AuthenticationError
    mod.BadRequestError = BadRequestError
    mod.APIConnectionError = APIConnectionError
    mod.Anthropic = lambda *a, **k: types.SimpleNamespace(
        messages=_FakeMessages(response, capture)
    )
    sys.modules["anthropic"] = mod
    return mod


def _run_with_fake(response, text="Build an app", skills=None, capture=None,
                   model_env=None):
    """Run parse_brief with a fake SDK + a dummy key; restore env/modules after.

    ``capture`` (a dict) receives the kwargs passed to ``messages.parse``.
    ``model_env`` sets ``ANTHROPIC_MODEL`` for the call (None → unset).
    """
    saved_mod = sys.modules.get("anthropic")
    saved_key = os.environ.get("ANTHROPIC_API_KEY")
    saved_model = os.environ.get("ANTHROPIC_MODEL")
    os.environ["ANTHROPIC_API_KEY"] = "sk-test-dummy"
    if model_env is not None:
        os.environ["ANTHROPIC_MODEL"] = model_env
    else:
        os.environ.pop("ANTHROPIC_MODEL", None)
    try:
        _install_fake_anthropic(response, capture)
        return parse_brief(text, skills if skills is not None else ["UX", "Research"])
    finally:
        if saved_mod is not None:
            sys.modules["anthropic"] = saved_mod
        else:
            sys.modules.pop("anthropic", None)
        if saved_key is None:
            os.environ.pop("ANTHROPIC_API_KEY", None)
        else:
            os.environ["ANTHROPIC_API_KEY"] = saved_key
        if saved_model is None:
            os.environ.pop("ANTHROPIC_MODEL", None)
        else:
            os.environ["ANTHROPIC_MODEL"] = saved_model


def test_parse_brief_unavailable_without_key():
    saved_key = os.environ.pop("ANTHROPIC_API_KEY", None)
    try:
        parse_brief("Build an app", ["UX"])
        assert False, "expected BriefParserUnavailable"
    except BriefParserUnavailable:
        pass
    finally:
        if saved_key is not None:
            os.environ["ANTHROPIC_API_KEY"] = saved_key


def test_parse_brief_happy_path_and_reconciles():
    parsed = _ModelOutput(
        draft_tasks=[
            DraftTask(task="Design UI", required_skill="ux", effort_hours=12),
            DraftTask(task="Smart contract", required_skill="Solidity", effort_hours=8),
        ],
        notes="Two tasks drafted.",
    )
    result = _run_with_fake(_FakeResponse(parsed), skills=["UX", "Research"])
    assert len(result.draft_tasks) == 2
    # Casing normalised against the vocabulary.
    assert result.draft_tasks[0].required_skill == "UX"
    # Out-of-vocab skill flagged + reported.
    assert result.unmatched_skills == ["Solidity"]
    assert result.draft_tasks[1].needs_user_review is True
    assert result.available_skills == ["UX", "Research"]
    assert result.notes == "Two tasks drafted."


def test_parse_brief_call_shape_matches_sdk():
    # The call must send output_format and must NOT send sampling params.
    parsed = _ModelOutput(draft_tasks=[], notes=None)
    capture = {}
    _run_with_fake(_FakeResponse(parsed), capture=capture)
    assert capture["output_format"] is _ModelOutput
    assert "messages" in capture and "max_tokens" in capture and "system" in capture
    for forbidden in ("temperature", "top_p", "top_k"):
        assert forbidden not in capture, f"{forbidden} should not be sent"


def test_parse_brief_uses_default_model():
    parsed = _ModelOutput(draft_tasks=[], notes=None)
    capture = {}
    _run_with_fake(_FakeResponse(parsed), capture=capture, model_env=None)
    assert capture["model"] == "claude-opus-4-8"


def test_parse_brief_model_overridable_by_env():
    parsed = _ModelOutput(draft_tasks=[], notes=None)
    capture = {}
    _run_with_fake(_FakeResponse(parsed), capture=capture, model_env="claude-haiku-4-5")
    assert capture["model"] == "claude-haiku-4-5"


def test_parse_brief_malformed_output_raises_clean_422():
    # Simulate messages.parse() raising pydantic.ValidationError on malformed
    # structured output (the real SDK validates via TypeAdapter.validate_json).
    from pydantic import ValidationError

    try:
        _ModelOutput.model_validate({"draft_tasks": [{}]})  # missing required fields
        assert False, "expected ValidationError to build the fixture"
    except ValidationError as built:
        validation_error = built

    try:
        _run_with_fake(validation_error)
        assert False, "expected BriefParserError"
    except BriefParserError as exc:
        assert exc.status == 422


def test_parse_brief_none_parsed_output_raises_502():
    try:
        _run_with_fake(_FakeResponse(None))
        assert False, "expected BriefParserError"
    except BriefParserError as exc:
        assert exc.status == 502


def test_parse_brief_refusal_raises_422():
    parsed = _ModelOutput(draft_tasks=[], notes=None)
    try:
        _run_with_fake(_FakeResponse(parsed, stop_reason="refusal"))
        assert False, "expected BriefParserError"
    except BriefParserError as exc:
        assert exc.status == 422


def test_parse_brief_rate_limit_raises_429():
    saved_mod = sys.modules.get("anthropic")
    saved_key = os.environ.get("ANTHROPIC_API_KEY")
    os.environ["ANTHROPIC_API_KEY"] = "sk-test-dummy"
    try:
        # Install the fake module, then make its client raise an error built
        # from THIS module's class (so parse_brief's except clause matches).
        mod = _install_fake_anthropic(None)
        err = mod.RateLimitError("busy")
        mod.Anthropic = lambda *a, **k: types.SimpleNamespace(
            messages=_FakeMessages(err)
        )
        parse_brief("Build an app", ["UX"])
        assert False, "expected BriefParserError"
    except BriefParserError as exc:
        assert exc.status == 429
    finally:
        if saved_mod is not None:
            sys.modules["anthropic"] = saved_mod
        else:
            sys.modules.pop("anthropic", None)
        if saved_key is None:
            os.environ.pop("ANTHROPIC_API_KEY", None)
        else:
            os.environ["ANTHROPIC_API_KEY"] = saved_key


# ---------------------------------------------------------------------------
# Endpoints
# ---------------------------------------------------------------------------

def test_extract_endpoint_docx_ok():
    content = _make_docx(["Launch a customer support portal with live chat."])
    r = client.post(
        "/projects/extract-brief-text",
        files={"file": ("brief.docx", io.BytesIO(content),
                        "application/vnd.openxmlformats-officedocument."
                        "wordprocessingml.document")},
    )
    assert r.status_code == 200
    body = r.json()
    assert body["file_type"] == "docx"
    assert "support portal" in body["text"]
    assert body["truncated"] is False


def test_extract_endpoint_rejects_txt():
    r = client.post(
        "/projects/extract-brief-text",
        files={"file": ("notes.txt", io.BytesIO(b"hello"), "text/plain")},
    )
    assert r.status_code == 415


def test_parse_brief_endpoint_503_without_key():
    saved_key = os.environ.pop("ANTHROPIC_API_KEY", None)
    try:
        r = client.post("/projects/parse-brief", json={"text": "Build an app"})
        assert r.status_code == 503
    finally:
        if saved_key is not None:
            os.environ["ANTHROPIC_API_KEY"] = saved_key


def test_parse_brief_endpoint_rejects_empty_text():
    r = client.post("/projects/parse-brief", json={"text": "   "})
    assert r.status_code == 422


def test_parse_brief_endpoint_ok_with_mocked_parser(monkeypatch=None):
    # Patch the route's parser to avoid any network; verify the route mapping.
    from src.api import routes
    from src.api.brief_parser import BriefParseResult

    def fake_parse(text, skills):
        assert text  # received the body text
        assert isinstance(skills, list) and skills  # real vocabulary injected
        return BriefParseResult(
            draft_tasks=[DraftTask(task="A", required_skill="UX", effort_hours=10)],
            available_skills=skills,
            unmatched_skills=[],
            notes="ok",
        )

    original = routes.brief_parser.parse_brief
    routes.brief_parser.parse_brief = fake_parse
    try:
        r = client.post("/projects/parse-brief", json={"text": "Build an app"})
        assert r.status_code == 200
        body = r.json()
        assert len(body["draft_tasks"]) == 1
        assert body["draft_tasks"][0]["task"] == "A"
        assert body["notes"] == "ok"
    finally:
        routes.brief_parser.parse_brief = original


def test_available_skills_matches_sample_data():
    from src.api import routes

    skills = routes._available_skills()
    # Drawn from the union of employee skills + AI-agent capabilities.
    assert "UX" in skills
    assert "Python" in skills
    assert skills == sorted(skills)


# ---------------------------------------------------------------------------
# Mapping: drafted tasks -> ProjectTaskInput accepted by /simulate/project
# ---------------------------------------------------------------------------

def test_drafted_tasks_map_into_project_simulation():
    drafts = [
        DraftTask(task="Research users", required_skill="Research", effort_hours=10,
                  needs_user_review=False),
        DraftTask(task="Build UI", required_skill="React", effort_hours=20,
                  dependencies=["Research users"]),
    ]
    # Mirror the frontend's toTaskInput mapping (UI-only flags dropped).
    tasks = [
        {
            "task": d.task,
            "required_skill": d.required_skill,
            "effort_hours": d.effort_hours,
            "priority": d.priority,
            "dependencies": d.dependencies,
            "is_required": True,
        }
        for d in drafts
    ]
    payload = {
        "tasks": tasks,
        "optimization_objective": "balanced",
        "current_team_human_names": ["Sarah", "Maya", "Alex"],
        "current_team_ai_agent_names": [],
    }
    r = client.post("/simulate/project", json=payload)
    assert r.status_code == 200
    body = r.json()
    assert len(body["options"]) == 6
    assert body["recommendation"]["recommended_option"] in body["options"]


# Allow running directly without pytest.
if __name__ == "__main__":
    failures = 0
    for name, fn in sorted(globals().items()):
        if name.startswith("test_") and callable(fn):
            try:
                fn()
                print(f"PASS  {name}")
            except AssertionError as exc:
                failures += 1
                print(f"FAIL  {name}: {exc}")
            except Exception as exc:  # noqa: BLE001
                failures += 1
                print(f"ERROR {name}: {exc!r}")
    if failures:
        print(f"\n{failures} test(s) failed.")
        sys.exit(1)
    print("\nAll brief tests passed.")

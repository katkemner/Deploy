"""Deterministic text extraction from an uploaded project brief.

This module is **LLM-free**. It turns an uploaded ``.docx`` or text-based
``.pdf`` into plain text so the text (never the raw file) can be previewed by
the user and, only after they confirm, sent to the AI-drafting step.

Design constraints honoured here:

* v1 supports ``.docx`` and **text-based** PDFs only — no OCR for scanned
  PDFs. A PDF with no extractable text is reported as a scanned/image-only
  error rather than silently returning an empty string.
* Nothing is persisted: callers pass the in-memory bytes; this module reads
  them and returns text. No file is written to disk.

The third-party parsers (``python-docx``, ``pypdf``) are imported lazily so the
rest of the API keeps working even if they are not installed.
"""

from __future__ import annotations

import io
from dataclasses import dataclass

# Reject absurdly large uploads before parsing (defensive; the route also caps).
MAX_BYTES = 5 * 1024 * 1024  # 5 MB
# Below this many non-whitespace characters a PDF is treated as scanned/empty.
MIN_TEXT_CHARS = 20
# Hard cap on text returned to the caller / forwarded to the LLM. Long briefs
# are truncated (with a flag) to keep the prompt bounded.
MAX_TEXT_CHARS = 50_000


class BriefExtractionError(Exception):
    """Raised on any problem extracting text. Carries an HTTP-ish ``status``.

    ``status`` mirrors the response the route should return (415 unsupported,
    413 too large, 422 unreadable / scanned). Keeping it here lets the route
    map the failure to a clean HTTP error without re-deriving the cause.
    """

    def __init__(self, status: int, message: str) -> None:
        super().__init__(message)
        self.status = status
        self.message = message


@dataclass
class ExtractedBrief:
    """Result of a successful extraction (no file is kept)."""

    filename: str
    file_type: str  # "docx" | "pdf"
    text: str
    char_count: int
    truncated: bool


def _detect_type(filename: str) -> str:
    """Return ``"docx"`` or ``"pdf"`` from the filename, else raise 415."""
    lower = (filename or "").lower()
    if lower.endswith(".docx"):
        return "docx"
    if lower.endswith(".pdf"):
        return "pdf"
    if lower.endswith(".doc"):
        raise BriefExtractionError(
            415,
            "Old-style .doc files aren't supported. Save as .docx (Word) or a "
            "text-based PDF and try again.",
        )
    raise BriefExtractionError(
        415, "Unsupported file type. Upload a Word (.docx) or a text-based PDF."
    )


def _extract_docx(content: bytes) -> str:
    """Extract paragraph + table-cell text from a .docx, in document order."""
    try:
        import docx  # python-docx
    except ImportError as exc:  # pragma: no cover - environment-dependent
        raise BriefExtractionError(
            503, "Document parsing isn't available on this server (python-docx)."
        ) from exc

    try:
        document = docx.Document(io.BytesIO(content))
    except Exception as exc:  # corrupt / not a real docx
        raise BriefExtractionError(
            422,
            "Couldn't read this Word file. It may be corrupted or password-"
            "protected.",
        ) from exc

    parts: list[str] = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_pdf(content: bytes) -> str:
    """Extract text from a text-based PDF (no OCR)."""
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - environment-dependent
        raise BriefExtractionError(
            503, "PDF parsing isn't available on this server (pypdf)."
        ) from exc

    try:
        reader = PdfReader(io.BytesIO(content))
    except Exception as exc:  # corrupt / not a real pdf
        raise BriefExtractionError(
            422,
            "Couldn't read this PDF. It may be corrupted or password-protected.",
        ) from exc

    pages: list[str] = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception:  # one bad page shouldn't fail the whole document
            continue
    return "\n".join(pages)


def extract_brief_text(content: bytes, filename: str) -> ExtractedBrief:
    """Extract plain text from an uploaded brief's bytes.

    Raises :class:`BriefExtractionError` (with an HTTP ``status``) on an empty
    upload, an unsupported type, an oversized file, an unreadable file, or a
    scanned/image-only PDF with no extractable text.
    """
    if not content:
        raise BriefExtractionError(422, "The uploaded file is empty.")
    if len(content) > MAX_BYTES:
        raise BriefExtractionError(
            413, "File exceeds the 5 MB limit. Upload a smaller brief."
        )

    file_type = _detect_type(filename)
    text = _extract_docx(content) if file_type == "docx" else _extract_pdf(content)
    text = (text or "").strip()

    # Scanned / image-only PDF (or an empty doc): no usable text.
    if len(text.replace(" ", "").replace("\n", "")) < MIN_TEXT_CHARS:
        if file_type == "pdf":
            raise BriefExtractionError(
                422,
                "This looks like a scanned or image-only PDF. v1 supports "
                "text-based PDFs only (no OCR). Paste the text or upload a "
                "Word (.docx) document instead.",
            )
        raise BriefExtractionError(
            422,
            "No readable text was found in this document. Check that it isn't "
            "empty and try again.",
        )

    truncated = len(text) > MAX_TEXT_CHARS
    if truncated:
        text = text[:MAX_TEXT_CHARS]

    return ExtractedBrief(
        filename=filename,
        file_type=file_type,
        text=text,
        char_count=len(text),
        truncated=truncated,
    )
